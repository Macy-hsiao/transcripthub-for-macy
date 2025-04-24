from flask import Flask, request, send_file
from pydub import AudioSegment
from pyannote.audio import Pipeline
from transformers import pipeline as hf_pipeline
from docx import Document
import os

app = Flask(__name__)

SPEECH_MODEL = "openai/whisper-large"
pipeline = hf_pipeline("automatic-speech-recognition", model=SPEECH_MODEL)
diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization", use_auth_token=os.getenv("HUGGINGFACE_TOKEN"))

@app.route('/transcribe', methods=['POST'])
def transcribe():
    file = request.files['file']
    file_path = "input_audio." + file.filename.split('.')[-1]
    file.save(file_path)
    
    diarization = diarization_pipeline(file_path)
    result = pipeline(file_path, return_timestamps=True)

    doc = Document()
    doc.add_heading('Transcript with Speaker Diarization', 0)
    
    for segment, _, speaker in diarization.itertracks(yield_label=True):
        start = int(segment.start)
        end = int(segment.end)
        text_segment = [s for s in result["chunks"] if s["timestamp"][0] >= start and s["timestamp"][0] <= end]
        text = " ".join([s["text"] for s in text_segment])
        doc.add_paragraph(f"[{start//3600:02}:{(start%3600)//60:02}:{start%60:02}] {speaker}: {text}")
    
    output_path = "transcript.docx"
    doc.save(output_path)
    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
